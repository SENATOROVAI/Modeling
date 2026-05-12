from pathlib import Path
import json
import shutil

import numpy as np
import scipy.io as sio
from scipy.interpolate import LinearNDInterpolator
from matplotlib import pyplot as plt
from matplotlib.patches import Rectangle
from matplotlib.tri import Triangulation
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/m/Documents/Цифровое моделирование")
ARTIFACTS = ROOT / "Артефакты_теплопроводность"
IMAGES = ARTIFACTS / "images"
DATA = ARTIFACTS / "data"
RENDERED = ARTIFACTS / "rendered"
DOCX_OUT = ROOT / "Otchet_teploprovodnost_chip.docx"

MINUS_MAT = ROOT / "MadeByStud" / "wrk_minus.mat"
PLUS_MAT = ROOT / "MadeByStud" / "wrk_plus.mat"

FIG_GEOMETRY = IMAGES / "01_geometry_processor.png"
FIG_MESH = IMAGES / "02_mesh.png"
FIG_FIELDS = IMAGES / "03_temperature_fields.png"
FIG_ISOTHERMS = IMAGES / "04_isotherms_range.png"
FIG_CHIP = IMAGES / "05_chip_options.png"
FIG_CODE = IMAGES / "06_code_fragment.png"
ALGORITHM_TXT = DATA / "chip_search_algorithm.txt"

ALGORITHM_LISTING = """minus = loadmat("wrk_minus.mat", squeeze_me=True)
plus = loadmat("wrk_plus.mat", squeeze_me=True)

p = minus["p1"]
t = minus["t1"][:3, :].astype(int) - 1
u1 = minus["u1"].ravel()
u2 = plus["u2"].ravel()

tri = Triangulation(p[0], p[1], t.T)
finder = tri.get_trifinder()
T1 = LinearNDInterpolator(np.c_[p[0], p[1]], u1)
T2 = LinearNDInterpolator(np.c_[p[0], p[1]], u2)

valid = []
for x0 in x_grid:
    for y0 in y_grid:
        x, y = chip_points(x0, y0, W, H)
        if np.any(finder(x, y) < 0):
            continue
        if intersects_processor(x0, y0, W, H):
            continue
        vals1 = T1(x, y)
        vals2 = T2(x, y)
        ok1 = vals1.min() >= -10 and vals1.max() <= 30
        ok2 = vals2.min() >= -10 and vals2.max() <= 30
        if ok1 and ok2:
            valid.append([x0, y0, vals1.mean(), vals2.mean()])

valid.sort(key=lambda row: row[3])
best_options = select_separated_positions(valid, count=2)"""

PARAMS = {
    "variant_source": "файлы MadeByStud/pde_Variant_0_minus.m, pde_Variant_0_plus.m и таблица на фото",
    "group": "ОТ-24-14",
    "discipline": "Цифровое моделирование",
    "topic": "Стационарная задача теплопроводности в MATLAB PDE Modeler",
    "shape_formula": "(E1 - R1) + P1",
    "R_m": 0.05,
    "R_cm": 5.0,
    "material": "гетинакс",
    "thermal_conductivity": 0.2149613069648,
    "insulated_boundary": 2,
    "processor_w_m": 0.015,
    "processor_h_m": 0.010,
    "processor_temperature": 40.0,
    "processor_bounds_m": (-0.0075, 0.0075, 0.0014, 0.0114),
    "chip_w_m": 0.010,
    "chip_h_m": 0.010,
    "chip_t_min": -10.0,
    "chip_t_max": 30.0,
    "external_t_min": -20.0,
    "external_t_max": 20.0,
}


def prepare_dirs():
    for folder in (ARTIFACTS, IMAGES, DATA, RENDERED):
        folder.mkdir(parents=True, exist_ok=True)


def cm_from_m(value):
    return value * 100


def load_solution():
    minus = sio.loadmat(MINUS_MAT, squeeze_me=True)
    plus = sio.loadmat(PLUS_MAT, squeeze_me=True)
    p = np.asarray(minus["p1"], dtype=float)
    e = np.asarray(minus["e1"], dtype=float)
    t = np.asarray(minus["t1"][:3, :], dtype=int) - 1
    u_minus = np.asarray(minus["u1"], dtype=float).ravel()
    u_plus = np.asarray(plus["u2"], dtype=float).ravel()
    return p, e, t, u_minus, u_plus


def board_boundary():
    r = PARAMS["R_m"]
    theta = np.linspace(0, np.pi, 180)
    top = np.column_stack([r * np.cos(theta), r * np.sin(theta)])
    left = np.array([[-r, 0], [0, -r]])
    right = np.array([[0, -r], [r, 0]])
    return np.vstack([top, left, right])


def centroid_board():
    r = PARAMS["R_m"]
    area_semi = 0.5 * np.pi * r * r
    area_tri = r * r
    y_semi = 4 * r / (3 * np.pi)
    y_tri = -r / 3
    yc = (area_semi * y_semi + area_tri * y_tri) / (area_semi + area_tri)
    return 0.0, yc, area_semi + area_tri


def chip_stats(interp, x0, y0, w, h, samples=17):
    sx = np.linspace(0, w, samples)
    sy = np.linspace(0, h, samples)
    dx, dy = np.meshgrid(sx, sy)
    vals = np.asarray(interp((x0 + dx).ravel(), (y0 + dy).ravel()), dtype=float)
    vals = vals[~np.isnan(vals)]
    return float(vals.min()), float(vals.max()), float(vals.mean())


def find_chip_options(p, t, u_minus, u_plus):
    tri = Triangulation(p[0], p[1], t.T)
    finder = tri.get_trifinder()
    interp_minus = LinearNDInterpolator(np.column_stack([p[0], p[1]]), u_minus)
    interp_plus = LinearNDInterpolator(np.column_stack([p[0], p[1]]), u_plus)

    w, h = PARAMS["chip_w_m"], PARAMS["chip_h_m"]
    tmin, tmax = PARAMS["chip_t_min"], PARAMS["chip_t_max"]
    proc = PARAMS["processor_bounds_m"]
    sx = np.linspace(0, w, 11)
    sy = np.linspace(0, h, 11)
    dx, dy = np.meshgrid(sx, sy)
    x_grid = np.arange(p[0].min(), p[0].max() - w + 1e-12, 0.0005)
    y_grid = np.arange(p[1].min(), p[1].max() - h + 1e-12, 0.0005)

    valid = []
    for x0 in x_grid:
        for y0 in y_grid:
            intersects_processor = not (x0 + w <= proc[0] or x0 >= proc[1] or y0 + h <= proc[2] or y0 >= proc[3])
            if intersects_processor:
                continue
            x = (x0 + dx).ravel()
            y = (y0 + dy).ravel()
            if np.any(finder(x, y) < 0):
                continue
            vals_minus = np.asarray(interp_minus(x, y), dtype=float)
            vals_plus = np.asarray(interp_plus(x, y), dtype=float)
            if np.any(np.isnan(vals_minus)) or np.any(np.isnan(vals_plus)):
                continue
            mn1, mx1, av1 = vals_minus.min(), vals_minus.max(), vals_minus.mean()
            mn2, mx2, av2 = vals_plus.min(), vals_plus.max(), vals_plus.mean()
            if mn1 >= tmin and mx1 <= tmax and mn2 >= tmin and mx2 <= tmax:
                valid.append(
                    {
                        "x0_m": float(x0),
                        "y0_m": float(y0),
                        "minus": (float(mn1), float(mx1), float(av1)),
                        "plus": (float(mn2), float(mx2), float(av2)),
                        "criterion_avg_warm": float(av2),
                        "worst_range": float(max(mx1 - mn1, mx2 - mn2)),
                    }
                )

    valid.sort(key=lambda item: (item["criterion_avg_warm"], item["worst_range"]))
    chosen = []
    for item in valid:
        if all(np.hypot(item["x0_m"] - old["x0_m"], item["y0_m"] - old["y0_m"]) > 0.02 for old in chosen):
            chosen.append(item)
        if len(chosen) == 2:
            break

    for item in chosen:
        item["minus"] = chip_stats(interp_minus, item["x0_m"], item["y0_m"], w, h)
        item["plus"] = chip_stats(interp_plus, item["x0_m"], item["y0_m"], w, h)
        item["criterion_avg_warm"] = item["plus"][2]
        item["worst_range"] = max(item["minus"][1] - item["minus"][0], item["plus"][1] - item["plus"][0])

    return tri, chosen, len(valid)


def add_processor(ax):
    x0, x1, y0, y1 = PARAMS["processor_bounds_m"]
    ax.add_patch(Rectangle((x0, y0), x1 - x0, y1 - y0, facecolor="white", edgecolor="black", linewidth=1.8))
    ax.text((x0 + x1) / 2, y1 + 0.002, "процессор", ha="center", va="bottom", fontsize=9, color="black")


def setup_axis(ax, title):
    ax.set_title(title, fontsize=11, color="black")
    ax.set_xlabel("x, м", color="black")
    ax.set_ylabel("y, м", color="black")
    ax.set_aspect("equal", adjustable="box")
    ax.tick_params(labelsize=8, colors="black")
    for spine in ax.spines.values():
        spine.set_color("black")


def save_fig(fig, path):
    fig.savefig(path, dpi=240, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def make_figures(p, t, u_minus, u_plus, tri, options):
    boundary = board_boundary()
    xc, yc, _ = centroid_board()

    fig, ax = plt.subplots(figsize=(6.3, 4.7), constrained_layout=True)
    ax.plot(boundary[:, 0], boundary[:, 1], color="black", linewidth=1.5)
    ax.fill(boundary[:, 0], boundary[:, 1], color="0.96")
    ax.plot(xc, yc, marker="o", markersize=5, color="black")
    ax.text(xc + 0.002, yc - 0.003, "центр тяжести", fontsize=9, color="black")
    add_processor(ax)
    setup_axis(ax, "Исходная область и положение процессора")
    save_fig(fig, FIG_GEOMETRY)

    fig, ax = plt.subplots(figsize=(6.3, 4.7), constrained_layout=True)
    ax.triplot(tri, color="0.55", linewidth=0.35)
    add_processor(ax)
    setup_axis(ax, "Конечноэлементная сетка расчетной области")
    save_fig(fig, FIG_MESH)

    vmin = min(float(u_minus.min()), float(u_plus.min()))
    vmax = max(float(u_minus.max()), float(u_plus.max()))
    levels = np.linspace(vmin, vmax, 18)
    fig, axes = plt.subplots(1, 2, figsize=(10.6, 4.6), constrained_layout=True)
    for ax, u, title in [
        (axes[0], u_minus, "Условие 1: внешняя температура -20 °C"),
        (axes[1], u_plus, "Условие 2: внешняя температура +20 °C"),
    ]:
        cf = ax.tricontourf(tri, u, levels=levels, cmap="Greys")
        ax.tricontour(tri, u, levels=[PARAMS["chip_t_min"], PARAMS["chip_t_max"]], colors="black", linewidths=0.8)
        add_processor(ax)
        setup_axis(ax, title)
        cbar = fig.colorbar(cf, ax=ax, shrink=0.86)
        cbar.ax.tick_params(colors="black", labelsize=8)
    save_fig(fig, FIG_FIELDS)

    fig, axes = plt.subplots(1, 2, figsize=(10.6, 4.6), constrained_layout=True)
    for ax, u, levels_i, title in [
        (axes[0], u_minus, [-10, 30], "Изотермы -10 °C и +30 °C при -20 °C"),
        (axes[1], u_plus, [30], "Изотерма +30 °C при +20 °C"),
    ]:
        ax.triplot(tri, color="0.85", linewidth=0.25)
        ax.tricontour(tri, u, levels=levels_i, colors="black", linewidths=1.0)
        add_processor(ax)
        setup_axis(ax, title)
    save_fig(fig, FIG_ISOTHERMS)

    fig, ax = plt.subplots(figsize=(6.5, 5.4), constrained_layout=True)
    ax.triplot(tri, color="0.84", linewidth=0.24)
    ax.tricontour(tri, u_plus, levels=np.linspace(20, 40, 11), colors="0.38", linewidths=0.7)
    add_processor(ax)
    for idx, item in enumerate(options, start=1):
        ax.add_patch(
            Rectangle(
                (item["x0_m"], item["y0_m"]),
                PARAMS["chip_w_m"],
                PARAMS["chip_h_m"],
                fill=False,
                edgecolor="black",
                linewidth=1.7,
                linestyle="--",
            )
        )
        ax.text(
            item["x0_m"] + PARAMS["chip_w_m"] / 2,
            item["y0_m"] + PARAMS["chip_h_m"] + 0.002,
            f"чип {idx}",
            ha="center",
            fontsize=9,
            color="black",
            bbox={"facecolor": "white", "edgecolor": "black", "linewidth": 0.5, "pad": 2},
        )
    setup_axis(ax, "Два допустимых варианта размещения чипа")
    save_fig(fig, FIG_CHIP)

    make_code_screenshot()


def make_code_screenshot():
    lines = ALGORITHM_LISTING.splitlines()
    font_candidates = [
        "/System/Library/Fonts/Menlo.ttc",
        "/System/Library/Fonts/Supplemental/Courier New.ttf",
        "/Library/Fonts/Courier New.ttf",
    ]
    font = None
    for candidate in font_candidates:
        if Path(candidate).exists():
            font = ImageFont.truetype(candidate, 28)
            break
    if font is None:
        font = ImageFont.load_default()
    line_height = 40
    margin_x = 34
    margin_y = 28
    width = 1500
    height = margin_y * 2 + line_height * len(lines)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width - 1, height - 1), outline="black", width=2)
    y = margin_y
    for line in lines:
        draw.text((margin_x, y), line, fill="black", font=font)
        y += line_height
    image.save(FIG_CODE)


def set_font_run(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold


def set_paragraph_format(p, first_line=True):
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_text(doc, text, size=14, bold=False, align=None, first_line=True):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font_run(run, size=size, bold=bold)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font_run(run, size=14, bold=True)
    return p


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font_run(run, size=14)
    return p


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")


def set_cell_text(cell, text, bold=False, size=14):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(str(text))
    set_font_run(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell)


def add_table(doc, rows, size=14):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell_text(table.cell(r, c), value, bold=(r == 0), size=size)
    return table


def add_picture(doc, path, caption, width_cm=15.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(0)
    cap.paragraph_format.line_spacing = 1.5
    run = cap.add_run(caption)
    set_font_run(run, size=14)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_font_run(run, size=14)


def configure_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_page_number(section)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)


def add_title_page(doc):
    add_text(doc, "ПЕРВОЕ ВЫСШЕЕ ТЕХНИЧЕСКОЕ УЧЕБНОЕ ЗАВЕДЕНИЕ РОССИИ", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "МИНИСТЕРСТВО науки и высшего ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "федеральное государственное бюджетное образовательное учреждение высшего образования", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Санкт-Петербургский горный университет императрицы Екатерины II", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Кафедра цифрового моделирования", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Отчет по практической работе", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "По дисциплине\t\tЦифровое моделирование", size=14, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "(наименование учебной дисциплины согласно учебному плану)", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Тема работы:\t\tСтационарная задача теплопроводности в MATLAB PDE Modeler", size=14, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "выполнил: студент гр.\t\tОТ-24-14\t\t\t\t________________", size=14, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "(шифр группы)\t\t(подпись)\t\t(Ф.И.О)", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Оценка: ____________________", size=14, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "Дата: ______________________", size=14, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "Проверил", size=14, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "руководитель работы:\t\tДоцент\t\t\t\tКосовцева Т.Р.", size=14, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "(должность)\t\t(подпись)\t\t(Ф.И.О.)", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Санкт-Петербург", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "2026", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    doc.add_page_break()


def build_report(options, valid_count):
    xc, yc, area = centroid_board()
    proc = PARAMS["processor_bounds_m"]
    best = min(options, key=lambda item: item["criterion_avg_warm"])
    second = [item for item in options if item is not best][0]

    doc = Document()
    configure_doc(doc)
    add_title_page(doc)

    add_heading(doc, "Введение")
    add_text(
        doc,
        "Стационарные задачи теплопроводности применяются при проектировании электронных устройств, когда необходимо заранее оценить температурный режим элементов на плате. "
        "В данной работе рассматривается плата контрольно-измерительного прибора из гетинакса. На плате требуется разместить процессор и микросхему памяти так, чтобы рабочая температура микросхемы оставалась в допустимом диапазоне при двух предельных температурах внешней среды.",
    )
    add_text(
        doc,
        "Расчет выполнен по данным задания в среде MATLAB PDE Modeler. Для дальнейшей обработки использованы сохраненные конечноэлементные решения из файлов wrk_minus.mat и wrk_plus.mat. "
        "Поиск положения микросхемы выполнен программно по сетке кандидатов с последующей проверкой температурного диапазона.",
    )

    add_heading(doc, "Практическая работа. СТАЦИОНАРНАЯ ЗАДАЧА ТЕПЛОПРОВОДНОСТИ")
    add_text(
        doc,
        "Цель: определить допустимую зону размещения чипа памяти на плате КИП и предложить два варианта компоновки. "
        "Из двух вариантов необходимо выбрать лучший по критерию минимальной средней температуры в зоне размещения чипа.",
    )

    add_heading(doc, "Исходные данные")
    add_table(
        doc,
        [
            ["Параметр", "Значение"],
            ["Материал платы", PARAMS["material"]],
            ["Форма расчетной области", PARAMS["shape_formula"]],
            ["Радиус платы R", f"{PARAMS['R_cm']:.1f} см"],
            ["Коэффициент теплопроводности", f"{PARAMS['thermal_conductivity']:.10f}"],
            ["Теплоизолированная граница", f"№ {PARAMS['insulated_boundary']}"],
            ["Температура внешней среды", "-20 °C и +20 °C"],
            ["Размер процессора", "1,5 x 1,0 см"],
            ["Температура процессора", "40 °C"],
            ["Размер чипа памяти", "1,0 x 1,0 см"],
            ["Допустимый диапазон чипа", "от -10 °C до +30 °C"],
        ],
    )

    add_heading(doc, "Расчетные соотношения и алгоритм")
    add_text(
        doc,
        "Температурное поле платы рассматривалось как стационарное. Поэтому в расчетной области решалось эллиптическое уравнение теплопроводности без внутренних источников тепла. "
        "Процессор учитывался как внутренняя граница с заданной температурой.",
    )
    add_formula(doc, "-div(k grad T) = 0")
    add_text(
        doc,
        "Граничные условия задавались следующим образом: на границе процессора температура равна 40 °C; на внешних неизолированных границах температура равна температуре внешней среды; "
        "на теплоизолированной границе тепловой поток равен нулю.",
    )
    add_formula(doc, "Tпроц = 40 °C;     Tвнеш = -20 °C или +20 °C;     q = -k*dT/dn = 0")
    add_text(
        doc,
        "Для проверки зоны размещения чипа в прямоугольнике вычислялись минимальная, максимальная и средняя температуры по контрольным точкам. "
        "Положение считалось допустимым только при выполнении температурного ограничения для обоих внешних режимов.",
    )
    add_formula(doc, "Tmin >= -10 °C;     Tmax <= +30 °C;     Tср = (T1 + T2 + ... + TN) / N")
    add_text(
        doc,
        "Ниже приведен укрупненный алгоритм программной проверки. Полный воспроизводимый скрипт находится в папке артефактов.",
    )
    add_picture(doc, FIG_CODE, "Рисунок 1. Фрагмент программы для поиска положения чипа", width_cm=15.3)

    add_heading(doc, "Построение области и размещение процессора")
    add_text(
        doc,
        "Расчетная область построена как объединение верхней полуокружности радиуса R и нижнего треугольника. "
        f"Площадь области составила {area * 10000:.2f} см2. По симметрии координата центра тяжести Xc = 0, расчетная координата Yc = {cm_from_m(yc):.2f} см.",
    )
    add_text(
        doc,
        f"Процессор размещен вблизи центра тяжести платы. Координаты его углов: x = [{cm_from_m(proc[0]):.2f}; {cm_from_m(proc[1]):.2f}] см, "
        f"y = [{cm_from_m(proc[2]):.2f}; {cm_from_m(proc[3]):.2f}] см. В модели процессор задан как внутренняя граница с температурой 40 °C.",
    )
    add_picture(doc, FIG_GEOMETRY, "Рисунок 2. Исходная область и положение процессора", width_cm=14.0)
    add_picture(doc, FIG_MESH, "Рисунок 3. Конечноэлементная сетка расчетной области", width_cm=14.0)

    add_heading(doc, "Граничные условия и расчет температурного поля")
    add_text(
        doc,
        "Для границы №2 задано условие теплоизоляции, то есть нулевой тепловой поток через границу. На остальных внешних границах задавалась температура внешней среды. "
        "Были рассмотрены два предельных режима: -20 °C и +20 °C. Температура процессора в обоих случаях оставалась равной 40 °C.",
    )
    add_text(
        doc,
        "Температурные поля построены в общей шкале, чтобы сравнение двух режимов было корректным. Более теплое внешнее условие является критичным для верхней границы допустимой температуры микросхемы.",
    )
    add_picture(doc, FIG_FIELDS, "Рисунок 4. Температурные поля для двух предельных внешних условий", width_cm=15.3)
    add_picture(doc, FIG_ISOTHERMS, "Рисунок 5. Изотермы допустимого диапазона температуры чипа", width_cm=15.3)

    add_heading(doc, "Поиск места для размещения чипа")
    add_text(
        doc,
        "Для проверки положения чипа использовался прямоугольник 1 x 1 см со сторонами, параллельными осям координат. "
        "Кандидат считался допустимым, если весь прямоугольник находится внутри платы, не пересекает процессор и во всех контрольных точках температура при обоих режимах лежит в диапазоне от -10 °C до +30 °C.",
    )
    add_text(
        doc,
        f"Всего программная проверка нашла {valid_count} допустимых положений на расчетной сетке. Для отчета выбраны два разнесенных варианта. "
        "Критерий выбора: минимальная средняя температура в зоне чипа при внешней температуре +20 °C.",
    )

    rows = [["Вариант", "Координаты углов чипа, см", "Tmin/Tmax/Tср при -20 °C", "Tmin/Tmax/Tср при +20 °C", "Критерий, °C"]]
    for idx, item in enumerate(options, start=1):
        x0, y0 = cm_from_m(item["x0_m"]), cm_from_m(item["y0_m"])
        x1, y1 = cm_from_m(item["x0_m"] + PARAMS["chip_w_m"]), cm_from_m(item["y0_m"] + PARAMS["chip_h_m"])
        mn, mx, av = item["minus"]
        pn, px, pa = item["plus"]
        rows.append(
            [
                str(idx),
                f"({x0:.2f}; {y0:.2f}) - ({x1:.2f}; {y1:.2f})",
                f"{mn:.2f} / {mx:.2f} / {av:.2f}",
                f"{pn:.2f} / {px:.2f} / {pa:.2f}",
                f"{item['criterion_avg_warm']:.2f}",
            ]
        )
    add_table(doc, rows, size=14)
    add_picture(doc, FIG_CHIP, "Рисунок 6. Два допустимых положения чипа на расчетной области", width_cm=14.4)

    add_heading(doc, "Выбор лучшего варианта")
    add_text(
        doc,
        f"Оба положения удовлетворяют температурному ограничению. Лучшим по принятому критерию является вариант 1: средняя температура в зоне чипа при внешней температуре +20 °C равна {best['plus'][2]:.2f} °C. "
        f"Для варианта 2 это значение равно {second['plus'][2]:.2f} °C. Разница небольшая, но первый вариант дает меньшую среднюю температуру в наиболее нагруженном тепловом режиме.",
    )

    add_heading(doc, "Заключение")
    add_text(
        doc,
        "В ходе работы была построена расчетная область платы, задано положение процессора, сформирована конечноэлементная сетка и выполнены два расчета стационарного температурного поля. "
        "Размещение чипа памяти при исходных условиях возможно. Найдены два допустимых варианта положения чипа, для каждого рассчитаны минимальная, максимальная и средняя температуры при внешних условиях -20 °C и +20 °C.",
    )
    add_text(
        doc,
        "По критерию минимальной средней температуры выбран первый вариант размещения. Выполнение работы позволило закрепить построение геометрии в PDE Modeler, задание граничных условий Дирихле и Неймана, анализ результатов конечноэлементного решения и программный поиск допустимой зоны размещения элемента.",
    )
    doc.save(DOCX_OUT)


def write_artifacts(options, valid_count):
    payload = {
        "parameters": PARAMS,
        "valid_positions_count": valid_count,
        "chosen_options": options,
    }
    (DATA / "results_summary.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    ALGORITHM_TXT.write_text(ALGORITHM_LISTING + "\n", encoding="utf-8")
    with (DATA / "results_summary.csv").open("w", encoding="utf-8") as f:
        f.write("variant,x0_cm,y0_cm,x1_cm,y1_cm,minus_min,minus_max,minus_avg,plus_min,plus_max,plus_avg,criterion_avg_warm\n")
        for idx, item in enumerate(options, start=1):
            f.write(
                f"{idx},{cm_from_m(item['x0_m']):.4f},{cm_from_m(item['y0_m']):.4f},"
                f"{cm_from_m(item['x0_m'] + PARAMS['chip_w_m']):.4f},{cm_from_m(item['y0_m'] + PARAMS['chip_h_m']):.4f},"
                f"{item['minus'][0]:.6f},{item['minus'][1]:.6f},{item['minus'][2]:.6f},"
                f"{item['plus'][0]:.6f},{item['plus'][1]:.6f},{item['plus'][2]:.6f},{item['criterion_avg_warm']:.6f}\n"
            )
    shutil.copy2(ROOT / "build_thermal_report.py", ARTIFACTS / "build_thermal_report.py")


def main():
    prepare_dirs()
    p, _e, t, u_minus, u_plus = load_solution()
    tri, options, valid_count = find_chip_options(p, t, u_minus, u_plus)
    make_figures(p, t, u_minus, u_plus, tri, options)
    write_artifacts(options, valid_count)
    build_report(options, valid_count)
    print(DOCX_OUT)
    print(ARTIFACTS)


if __name__ == "__main__":
    main()
