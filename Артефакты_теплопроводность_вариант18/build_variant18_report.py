from pathlib import Path
import json
import shutil

import numpy as np
from scipy.interpolate import LinearNDInterpolator
from scipy.spatial import Delaunay
from scipy import sparse
from scipy.sparse.linalg import spsolve
from matplotlib import pyplot as plt
from matplotlib.patches import Rectangle
from matplotlib.path import Path as MplPath
from matplotlib.tri import Triangulation
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/m/Documents/Цифровое моделирование")
ARTIFACTS = ROOT / "Артефакты_теплопроводность_вариант18"
IMAGES = ARTIFACTS / "images"
DATA = ARTIFACTS / "data"
RENDERED = ARTIFACTS / "rendered"
DOCX_OUT = ROOT / "Otchet_teploprovodnost_chip_variant18.docx"

FIG_GEOMETRY = IMAGES / "01_geometry_processor.png"
FIG_MESH = IMAGES / "02_mesh.png"
FIG_FIELDS = IMAGES / "03_temperature_fields.png"
FIG_ISOTHERMS = IMAGES / "04_isotherms_range.png"
FIG_CHIP = IMAGES / "05_chip_options.png"
FIG_CODE = IMAGES / "06_code_fragment.png"

PARAMS = {
    "variant": 18,
    "group": "ОТ-24-14",
    "shape": "область 18",
    "R_m": 0.06,
    "R_cm": 6.0,
    "material": "гетинакс",
    "thermal_conductivity": 0.2149613069648,
    "insulated_boundary": 2,
    "processor_w_m": 0.010,
    "processor_h_m": 0.015,
    "processor_temperature": 40.0,
    "chip_w_m": 0.010,
    "chip_h_m": 0.010,
    "chip_t_min": -10.0,
    "chip_t_max": 30.0,
    "external_t_min": -20.0,
    "external_t_max": 20.0,
}

CODE_FRAGMENT = """points, triangles, boundary = build_mesh_variant18(R)
K = assemble_stiffness(points, triangles, k)

for Text in [-20, 20]:
    dirichlet = external_boundary_nodes(boundary, except_label=2)
    dirichlet += processor_nodes(points, processor_rect)
    values = np.full(len(points), Text)
    values[processor_nodes] = 40
    T = solve_laplace(K, dirichlet, values)
    fields.append(T)

valid = []
for x0 in x_grid:
    for y0 in y_grid:
        x, y = chip_points(x0, y0, Wmicro, Hmicro)
        if inside_board(x, y) and not intersects_processor(x0, y0):
            t1 = interp_minus(x, y)
            t2 = interp_plus(x, y)
            if t1.min() >= -10 and t1.max() <= 30:
                if t2.min() >= -10 and t2.max() <= 30:
                    valid.append([x0, y0, t1.mean(), t2.mean()])

valid.sort(key=lambda row: row[3])
best_options = select_separated_positions(valid, count=2)"""


def prepare_dirs():
    for folder in (ARTIFACTS, IMAGES, DATA, RENDERED):
        folder.mkdir(parents=True, exist_ok=True)


def cm_from_m(value):
    return value * 100


def boundary_points_with_labels(r, n_arc=45, n_seg=30):
    pts = []
    labels = []

    def add_arc(a0, a1, label, n=n_arc, skip_first=False):
        angles = np.linspace(a0, a1, n)
        if skip_first:
            angles = angles[1:]
        for a in angles:
            pts.append((r * np.cos(a), r * np.sin(a)))
            labels.append(label)

    def add_seg(p0, p1, label, n=n_seg, skip_first=True):
        ts = np.linspace(0, 1, n)
        if skip_first:
            ts = ts[1:]
        p0 = np.array(p0, dtype=float)
        p1 = np.array(p1, dtype=float)
        for t in ts:
            p = p0 * (1 - t) + p1 * t
            pts.append(tuple(p))
            labels.append(label)

    p_left = (-r, 0.0)
    p_top = (0.0, r)
    p_right = (r, 0.0)
    p_center = (0.0, 0.0)
    p_diag = (r / np.sqrt(2), -r / np.sqrt(2))
    p_bottom = (0.0, -r)

    add_arc(np.pi, np.pi / 2, 1, skip_first=False)
    add_arc(np.pi / 2, 0.0, 2, skip_first=True)
    add_seg(p_right, p_center, 3, skip_first=True)
    add_seg(p_center, p_diag, 4, skip_first=True)
    add_arc(-np.pi / 4, -np.pi / 2, 5, skip_first=True)
    add_arc(-np.pi / 2, -np.pi, 6, skip_first=True)
    return np.array(pts, dtype=float), np.array(labels, dtype=int)


def board_path():
    pts, _ = boundary_points_with_labels(PARAMS["R_m"], n_arc=80, n_seg=45)
    return MplPath(pts, closed=True), pts


def polygon_centroid(poly):
    x = poly[:, 0]
    y = poly[:, 1]
    x2 = np.r_[x, x[0]]
    y2 = np.r_[y, y[0]]
    cross = x2[:-1] * y2[1:] - x2[1:] * y2[:-1]
    area = 0.5 * cross.sum()
    cx = ((x2[:-1] + x2[1:]) * cross).sum() / (6 * area)
    cy = ((y2[:-1] + y2[1:]) * cross).sum() / (6 * area)
    return float(cx), float(cy), abs(float(area))


def processor_bounds(cx, cy):
    w = PARAMS["processor_w_m"]
    h = PARAMS["processor_h_m"]
    return (cx - w / 2, cx + w / 2, cy - h / 2, cy + h / 2)


def rect_points(bounds, n=8):
    x0, x1, y0, y1 = bounds
    pts = []
    for x in np.linspace(x0, x1, n):
        pts.append((x, y0))
        pts.append((x, y1))
    for y in np.linspace(y0, y1, n):
        pts.append((x0, y))
        pts.append((x1, y))
    return np.array(pts, dtype=float)


def build_mesh():
    r = PARAMS["R_m"]
    path, poly = board_path()
    cx, cy, area = polygon_centroid(poly)
    proc = processor_bounds(cx, cy)

    bpts, blabels = boundary_points_with_labels(r)
    h = 0.0025
    xs = np.arange(-r, r + h / 2, h)
    ys = np.arange(-r, r + h / 2, h)
    xx, yy = np.meshgrid(xs, ys)
    grid = np.column_stack([xx.ravel(), yy.ravel()])
    inside = path.contains_points(grid, radius=1e-10)
    grid = grid[inside]
    p_rect = rect_points(proc, n=9)
    points_raw = np.vstack([bpts, grid, p_rect])
    rounded = np.round(points_raw, 8)
    _, unique_idx = np.unique(rounded, axis=0, return_index=True)
    points = points_raw[np.sort(unique_idx)]

    deln = Delaunay(points)
    tris = deln.simplices
    centers = points[tris].mean(axis=1)
    keep = path.contains_points(centers, radius=1e-10)
    triangles = tris[keep]

    boundary_nodes = {}
    for label in range(1, 7):
        label_pts = bpts[blabels == label]
        node_ids = []
        for bp in label_pts:
            dist = np.linalg.norm(points - bp, axis=1)
            node_ids.append(int(np.argmin(dist)))
        boundary_nodes[label] = sorted(set(node_ids))

    proc_nodes = np.where(
        (points[:, 0] >= proc[0] - 1e-12)
        & (points[:, 0] <= proc[1] + 1e-12)
        & (points[:, 1] >= proc[2] - 1e-12)
        & (points[:, 1] <= proc[3] + 1e-12)
    )[0]
    return points, triangles, boundary_nodes, proc_nodes, proc, path, (cx, cy, area)


def assemble_stiffness(points, triangles):
    rows = []
    cols = []
    data = []
    k = PARAMS["thermal_conductivity"]
    for tri in triangles:
        coords = points[tri]
        x = coords[:, 0]
        y = coords[:, 1]
        area = 0.5 * abs((x[1] - x[0]) * (y[2] - y[0]) - (x[2] - x[0]) * (y[1] - y[0]))
        if area <= 1e-14:
            continue
        b = np.array([y[1] - y[2], y[2] - y[0], y[0] - y[1]])
        c = np.array([x[2] - x[1], x[0] - x[2], x[1] - x[0]])
        ke = k * (np.outer(b, b) + np.outer(c, c)) / (4 * area)
        for i in range(3):
            for j in range(3):
                rows.append(tri[i])
                cols.append(tri[j])
                data.append(ke[i, j])
    return sparse.csr_matrix((data, (rows, cols)), shape=(len(points), len(points)))


def solve_field(K, points, boundary_nodes, proc_nodes, text):
    n = len(points)
    dirichlet = set()
    for label, nodes in boundary_nodes.items():
        if label != PARAMS["insulated_boundary"]:
            dirichlet.update(nodes)
    dirichlet.update(map(int, proc_nodes))
    dirichlet = np.array(sorted(dirichlet), dtype=int)
    fixed = np.zeros(n, dtype=bool)
    fixed[dirichlet] = True
    free = np.where(~fixed)[0]
    values = np.zeros(n)
    values[dirichlet] = text
    values[np.intersect1d(dirichlet, proc_nodes)] = PARAMS["processor_temperature"]
    rhs = -K[free][:, dirichlet] @ values[dirichlet]
    values[free] = spsolve(K[free][:, free], rhs)
    return values


def chip_temperature_stats(interp, x0, y0, w, h, samples=17):
    sx = np.linspace(0, w, samples)
    sy = np.linspace(0, h, samples)
    dx, dy = np.meshgrid(sx, sy)
    vals = np.asarray(interp((x0 + dx).ravel(), (y0 + dy).ravel()), dtype=float)
    vals = vals[~np.isnan(vals)]
    return float(vals.min()), float(vals.max()), float(vals.mean())


def find_chip_options(points, triangles, u_minus, u_plus, proc, path):
    interp_minus = LinearNDInterpolator(points, u_minus)
    interp_plus = LinearNDInterpolator(points, u_plus)
    w = PARAMS["chip_w_m"]
    h = PARAMS["chip_h_m"]
    tmin = PARAMS["chip_t_min"]
    tmax = PARAMS["chip_t_max"]
    r = PARAMS["R_m"]
    sx = np.linspace(0, w, 11)
    sy = np.linspace(0, h, 11)
    dx, dy = np.meshgrid(sx, sy)
    valid = []
    for x0 in np.arange(-r, r - w + 1e-12, 0.001):
        for y0 in np.arange(-r, r - h + 1e-12, 0.001):
            if not (x0 + w <= proc[0] or x0 >= proc[1] or y0 + h <= proc[2] or y0 >= proc[3]):
                continue
            x = (x0 + dx).ravel()
            y = (y0 + dy).ravel()
            if not np.all(path.contains_points(np.column_stack([x, y]), radius=1e-10)):
                continue
            vals1 = np.asarray(interp_minus(x, y), dtype=float)
            vals2 = np.asarray(interp_plus(x, y), dtype=float)
            if np.any(np.isnan(vals1)) or np.any(np.isnan(vals2)):
                continue
            mn1, mx1, av1 = vals1.min(), vals1.max(), vals1.mean()
            mn2, mx2, av2 = vals2.min(), vals2.max(), vals2.mean()
            if mn1 >= tmin and mx1 <= tmax and mn2 >= tmin and mx2 <= tmax:
                valid.append(
                    {
                        "x0_m": float(x0),
                        "y0_m": float(y0),
                        "minus": (float(mn1), float(mx1), float(av1)),
                        "plus": (float(mn2), float(mx2), float(av2)),
                        "criterion_avg_warm": float(av2),
                    }
                )
    valid.sort(key=lambda item: (item["criterion_avg_warm"], item["plus"][1] - item["plus"][0]))
    chosen = []
    for item in valid:
        if all(np.hypot(item["x0_m"] - old["x0_m"], item["y0_m"] - old["y0_m"]) > 0.02 for old in chosen):
            chosen.append(item)
        if len(chosen) == 2:
            break
    for item in chosen:
        item["minus"] = chip_temperature_stats(interp_minus, item["x0_m"], item["y0_m"], w, h)
        item["plus"] = chip_temperature_stats(interp_plus, item["x0_m"], item["y0_m"], w, h)
        item["criterion_avg_warm"] = item["plus"][2]
    return chosen, len(valid)


def make_code_image():
    lines = CODE_FRAGMENT.splitlines()
    font_paths = [
        "/System/Library/Fonts/Menlo.ttc",
        "/System/Library/Fonts/Supplemental/Courier New.ttf",
        "/Library/Fonts/Courier New.ttf",
    ]
    font = None
    for path in font_paths:
        if Path(path).exists():
            font = ImageFont.truetype(path, 26)
            break
    if font is None:
        font = ImageFont.load_default()
    width = 1500
    line_height = 38
    margin_x = 34
    margin_y = 28
    height = margin_y * 2 + line_height * len(lines)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, width - 1, height - 1), outline="black", width=2)
    for i, line in enumerate(lines):
        draw.text((margin_x, margin_y + i * line_height), line, fill="black", font=font)
    image.save(FIG_CODE)


def setup_axis(ax, title):
    ax.set_title(title, fontsize=11, color="black")
    ax.set_xlabel("x, м", color="black")
    ax.set_ylabel("y, м", color="black")
    ax.set_aspect("equal", adjustable="box")
    ax.tick_params(labelsize=8, colors="black")
    for spine in ax.spines.values():
        spine.set_color("black")


def add_processor(ax, proc):
    x0, x1, y0, y1 = proc
    ax.add_patch(Rectangle((x0, y0), x1 - x0, y1 - y0, facecolor="white", edgecolor="black", linewidth=1.7))
    ax.text((x0 + x1) / 2, y1 + 0.002, "процессор", ha="center", fontsize=9, color="black")


def save_fig(fig, path):
    fig.savefig(path, dpi=240, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def make_figures(points, triangles, boundary_nodes, proc, centroid, u_minus, u_plus, options):
    tri = Triangulation(points[:, 0], points[:, 1], triangles)
    bpts, _ = boundary_points_with_labels(PARAMS["R_m"], n_arc=90, n_seg=50)
    make_code_image()

    fig, ax = plt.subplots(figsize=(6.3, 4.8), constrained_layout=True)
    ax.plot(bpts[:, 0], bpts[:, 1], color="black", linewidth=1.6)
    ax.fill(bpts[:, 0], bpts[:, 1], color="0.96")
    ax.plot(centroid[0], centroid[1], "ko", markersize=5)
    ax.text(centroid[0] + 0.002, centroid[1] - 0.003, "центр тяжести", fontsize=9, color="black")
    add_processor(ax, proc)
    setup_axis(ax, "Область 18 и положение процессора")
    save_fig(fig, FIG_GEOMETRY)

    fig, ax = plt.subplots(figsize=(6.3, 4.8), constrained_layout=True)
    ax.triplot(tri, color="0.55", linewidth=0.3)
    add_processor(ax, proc)
    setup_axis(ax, "Конечноэлементная сетка области 18")
    save_fig(fig, FIG_MESH)

    levels = np.linspace(min(u_minus.min(), u_plus.min()), max(u_minus.max(), u_plus.max()), 18)
    fig, axes = plt.subplots(1, 2, figsize=(10.6, 4.6), constrained_layout=True)
    for ax, u, title in [
        (axes[0], u_minus, "Условие 1: внешняя температура -20 °C"),
        (axes[1], u_plus, "Условие 2: внешняя температура +20 °C"),
    ]:
        cf = ax.tricontourf(tri, u, levels=levels, cmap="Greys")
        ax.tricontour(tri, u, levels=[PARAMS["chip_t_min"], PARAMS["chip_t_max"]], colors="black", linewidths=0.8)
        add_processor(ax, proc)
        setup_axis(ax, title)
        cbar = fig.colorbar(cf, ax=ax, shrink=0.86)
        cbar.ax.tick_params(colors="black", labelsize=8)
    save_fig(fig, FIG_FIELDS)

    fig, axes = plt.subplots(1, 2, figsize=(10.6, 4.6), constrained_layout=True)
    for ax, u, lvls, title in [
        (axes[0], u_minus, [-10, 30], "Изотермы -10 °C и +30 °C при -20 °C"),
        (axes[1], u_plus, [30], "Изотерма +30 °C при +20 °C"),
    ]:
        ax.triplot(tri, color="0.86", linewidth=0.22)
        ax.tricontour(tri, u, levels=lvls, colors="black", linewidths=1.0)
        add_processor(ax, proc)
        setup_axis(ax, title)
    save_fig(fig, FIG_ISOTHERMS)

    fig, ax = plt.subplots(figsize=(6.5, 5.3), constrained_layout=True)
    ax.triplot(tri, color="0.86", linewidth=0.22)
    ax.tricontour(tri, u_plus, levels=np.linspace(20, 40, 11), colors="0.38", linewidths=0.7)
    add_processor(ax, proc)
    for idx, item in enumerate(options, start=1):
        ax.add_patch(
            Rectangle(
                (item["x0_m"], item["y0_m"]),
                PARAMS["chip_w_m"],
                PARAMS["chip_h_m"],
                fill=False,
                edgecolor="black",
                linewidth=1.7,
                linestyle="--",
            )
        )
        ax.text(
            item["x0_m"] + PARAMS["chip_w_m"] / 2,
            item["y0_m"] + PARAMS["chip_h_m"] + 0.002,
            f"чип {idx}",
            ha="center",
            fontsize=9,
            color="black",
            bbox={"facecolor": "white", "edgecolor": "black", "linewidth": 0.5, "pad": 2},
        )
    setup_axis(ax, "Два допустимых варианта размещения чипа")
    save_fig(fig, FIG_CHIP)


def set_font_run(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold


def set_paragraph_format(p, first_line=True):
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_text(doc, text, size=14, bold=False, align=None, first_line=True):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font_run(run, size=size, bold=bold)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font_run(run, size=14, bold=True)
    return p


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font_run(run, size=14)


def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "000000")


def set_cell_text(cell, text, bold=False, size=14):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(str(text))
    set_font_run(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell)


def add_table(doc, rows, size=14):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            set_cell_text(table.cell(r, c), value, bold=(r == 0), size=size)
    return table


def add_picture(doc, path, caption, width_cm=15.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(0)
    cap.paragraph_format.line_spacing = 1.5
    run = cap.add_run(caption)
    set_font_run(run, size=14)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_font_run(run, size=14)


def configure_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_page_number(section)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)


def add_title_page(doc):
    add_text(doc, "ПЕРВОЕ ВЫСШЕЕ ТЕХНИЧЕСКОЕ УЧЕБНОЕ ЗАВЕДЕНИЕ РОССИИ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "МИНИСТЕРСТВО науки и высшего ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "федеральное государственное бюджетное образовательное учреждение высшего образования", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Санкт-Петербургский горный университет императрицы Екатерины II", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Кафедра цифрового моделирования", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Отчет по практической работе", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "По дисциплине\t\tЦифровое моделирование", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "(наименование учебной дисциплины согласно учебному плану)", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Тема работы:\t\tСтационарная задача теплопроводности в MATLAB PDE Modeler", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "Вариант:\t\t18", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "выполнил: студент гр.\t\tОТ-24-14\t\t\t\t________________", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "(шифр группы)\t\t(подпись)\t\t(Ф.И.О)", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Оценка: ____________________", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "Дата: ______________________", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "Проверил", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "руководитель работы:\t\tДоцент\t\t\t\tКосовцева Т.Р.", align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
    add_text(doc, "(должность)\t\t(подпись)\t\t(Ф.И.О.)", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "Санкт-Петербург", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_text(doc, "2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    doc.add_page_break()


def build_report(options, valid_count, centroid, area, proc):
    best = min(options, key=lambda item: item["criterion_avg_warm"])
    second = [item for item in options if item is not best][0]
    doc = Document()
    configure_doc(doc)
    add_title_page(doc)
    add_heading(doc, "Введение")
    add_text(
        doc,
        "Стационарные задачи теплопроводности применяются при проектировании электронных устройств, когда необходимо заранее оценить температурный режим элементов на плате. В данной работе рассматривается вариант 18: плата контрольно-измерительного прибора из гетинакса, на которой требуется разместить процессор и микросхему памяти.",
    )
    add_text(
        doc,
        "Расчет выполнен по параметрам из таблицы вариантов для группы ОТ-24-14. Для варианта 18 построена отдельная расчетная область, задана теплоизолированная граница №2 и выполнены два расчета для внешних температур -20 °C и +20 °C.",
    )
    add_heading(doc, "Практическая работа. СТАЦИОНАРНАЯ ЗАДАЧА ТЕПЛОПРОВОДНОСТИ")
    add_text(
        doc,
        "Цель: определить допустимую зону размещения чипа памяти на плате КИП и предложить два варианта компоновки. Из двух вариантов необходимо выбрать лучший по критерию минимальной средней температуры в зоне размещения чипа.",
    )
    add_heading(doc, "Исходные данные")
    add_table(
        doc,
        [
            ["Параметр", "Значение"],
            ["Вариант", "18"],
            ["Форма платы", "область 18"],
            ["Материал платы", PARAMS["material"]],
            ["Радиус платы R", "6,0 см"],
            ["Теплоизолированная граница", "№2"],
            ["Температура внешней среды", "-20 °C и +20 °C"],
            ["Размер процессора", "1,0 x 1,5 см"],
            ["Температура процессора", "40 °C"],
            ["Размер чипа памяти", "1,0 x 1,0 см"],
            ["Допустимый диапазон чипа", "от -10 °C до +30 °C"],
        ],
    )
    add_heading(doc, "Расчетные соотношения и алгоритм")
    add_text(
        doc,
        "Температурное поле платы рассматривалось как стационарное. В расчетной области решалось уравнение теплопроводности без внутренних источников тепла. Процессор учитывался как область с заданной температурой.",
    )
    add_formula(doc, "-div(k grad T) = 0")
    add_text(
        doc,
        "На границе процессора задана температура 40 °C. На внешних неизолированных границах задана температура внешней среды. На границе №2 задано условие теплоизоляции, то есть нулевой тепловой поток.",
    )
    add_formula(doc, "Tпроц = 40 °C;     Tвнеш = -20 °C или +20 °C;     q = -k*dT/dn = 0")
    add_text(
        doc,
        "Для проверки положения чипа вычислялись минимальная, максимальная и средняя температуры в прямоугольной зоне 1 x 1 см. Положение считалось допустимым, если температурное ограничение выполнялось при обоих внешних режимах.",
    )
    add_formula(doc, "Tmin >= -10 °C;     Tmax <= +30 °C;     Tср = (T1 + T2 + ... + TN) / N")
    add_picture(doc, FIG_CODE, "Рисунок 1. Фрагмент программы для поиска положения чипа", width_cm=15.3)
    add_heading(doc, "Построение области и размещение процессора")
    add_text(
        doc,
        f"Расчетная область соответствует рисунку для области 18. Площадь области по построенной геометрии составила {area * 10000:.2f} см2. Координаты центра тяжести: Xc = {cm_from_m(centroid[0]):.2f} см, Yc = {cm_from_m(centroid[1]):.2f} см.",
    )
    add_text(
        doc,
        f"Процессор размещен центром в центре тяжести платы. Координаты его углов: x = [{cm_from_m(proc[0]):.2f}; {cm_from_m(proc[1]):.2f}] см, y = [{cm_from_m(proc[2]):.2f}; {cm_from_m(proc[3]):.2f}] см.",
    )
    add_picture(doc, FIG_GEOMETRY, "Рисунок 2. Область 18 и положение процессора", width_cm=14.0)
    add_picture(doc, FIG_MESH, "Рисунок 3. Конечноэлементная сетка области 18", width_cm=14.0)
    add_heading(doc, "Граничные условия и расчет температурного поля")
    add_text(
        doc,
        "Для границы №2 задано условие теплоизоляции. На остальных внешних границах задавалась температура внешней среды. Были рассмотрены два предельных режима: -20 °C и +20 °C. Температура процессора в обоих случаях оставалась равной 40 °C.",
    )
    add_text(
        doc,
        "Температурные поля построены в общей шкале, чтобы сравнение двух режимов было корректным. Более теплое внешнее условие является критичным для верхней границы допустимой температуры микросхемы.",
    )
    add_picture(doc, FIG_FIELDS, "Рисунок 4. Температурные поля для двух предельных внешних условий", width_cm=15.3)
    add_picture(doc, FIG_ISOTHERMS, "Рисунок 5. Изотермы допустимого диапазона температуры чипа", width_cm=15.3)
    add_heading(doc, "Поиск места для размещения чипа")
    add_text(
        doc,
        "Для проверки положения чипа использовался прямоугольник 1 x 1 см со сторонами, параллельными осям координат. Кандидат считался допустимым, если весь прямоугольник находится внутри платы, не пересекает процессор и во всех контрольных точках температура при обоих режимах лежит в диапазоне от -10 °C до +30 °C.",
    )
    add_text(
        doc,
        f"Всего программная проверка нашла {valid_count} допустимых положений на расчетной сетке. Для отчета выбраны два разнесенных варианта. Критерий выбора: минимальная средняя температура в зоне чипа при внешней температуре +20 °C.",
    )
    rows = [["Вариант", "Координаты углов чипа, см", "Tmin/Tmax/Tср при -20 °C", "Tmin/Tmax/Tср при +20 °C", "Критерий, °C"]]
    for idx, item in enumerate(options, start=1):
        x0, y0 = cm_from_m(item["x0_m"]), cm_from_m(item["y0_m"])
        x1, y1 = cm_from_m(item["x0_m"] + PARAMS["chip_w_m"]), cm_from_m(item["y0_m"] + PARAMS["chip_h_m"])
        mn, mx, av = item["minus"]
        pn, px, pa = item["plus"]
        rows.append([str(idx), f"({x0:.2f}; {y0:.2f}) - ({x1:.2f}; {y1:.2f})", f"{mn:.2f} / {mx:.2f} / {av:.2f}", f"{pn:.2f} / {px:.2f} / {pa:.2f}", f"{item['criterion_avg_warm']:.2f}"])
    add_table(doc, rows)
    add_picture(doc, FIG_CHIP, "Рисунок 6. Два допустимых положения чипа на расчетной области", width_cm=14.4)
    add_heading(doc, "Выбор лучшего варианта")
    add_text(
        doc,
        f"Оба положения удовлетворяют температурному ограничению. Лучшим по принятому критерию является вариант 1: средняя температура в зоне чипа при внешней температуре +20 °C равна {best['plus'][2]:.2f} °C. Для варианта 2 это значение равно {second['plus'][2]:.2f} °C.",
    )
    add_heading(doc, "Заключение")
    add_text(
        doc,
        "В ходе работы была построена расчетная область для варианта 18, сформирована конечноэлементная сетка и выполнены два расчета стационарного температурного поля. Размещение чипа памяти при исходных условиях возможно. Найдены два допустимых варианта положения чипа, для каждого рассчитаны минимальная, максимальная и средняя температуры при внешних условиях -20 °C и +20 °C.",
    )
    add_text(
        doc,
        "По критерию минимальной средней температуры выбран первый вариант размещения. Выполнение работы позволило закрепить построение геометрии, задание граничных условий Дирихле и Неймана, анализ результатов конечноэлементного решения и программный поиск допустимой зоны размещения элемента.",
    )
    doc.save(DOCX_OUT)


def write_artifacts(options, valid_count, centroid, area, proc):
    payload = {
        "parameters": PARAMS,
        "centroid_m": centroid,
        "area_m2": area,
        "processor_bounds_m": proc,
        "valid_positions_count": valid_count,
        "chosen_options": options,
    }
    (DATA / "results_summary.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    (DATA / "chip_search_code_fragment.txt").write_text(CODE_FRAGMENT + "\n", encoding="utf-8")
    with (DATA / "results_summary.csv").open("w", encoding="utf-8") as f:
        f.write("variant,x0_cm,y0_cm,x1_cm,y1_cm,minus_min,minus_max,minus_avg,plus_min,plus_max,plus_avg,criterion_avg_warm\n")
        for idx, item in enumerate(options, start=1):
            f.write(
                f"{idx},{cm_from_m(item['x0_m']):.4f},{cm_from_m(item['y0_m']):.4f},"
                f"{cm_from_m(item['x0_m'] + PARAMS['chip_w_m']):.4f},{cm_from_m(item['y0_m'] + PARAMS['chip_h_m']):.4f},"
                f"{item['minus'][0]:.6f},{item['minus'][1]:.6f},{item['minus'][2]:.6f},"
                f"{item['plus'][0]:.6f},{item['plus'][1]:.6f},{item['plus'][2]:.6f},{item['criterion_avg_warm']:.6f}\n"
            )
    shutil.copy2(ROOT / "build_variant18_report.py", ARTIFACTS / "build_variant18_report.py")


def main():
    prepare_dirs()
    points, triangles, boundary_nodes, proc_nodes, proc, path, centroid_info = build_mesh()
    K = assemble_stiffness(points, triangles)
    u_minus = solve_field(K, points, boundary_nodes, proc_nodes, PARAMS["external_t_min"])
    u_plus = solve_field(K, points, boundary_nodes, proc_nodes, PARAMS["external_t_max"])
    options, valid_count = find_chip_options(points, triangles, u_minus, u_plus, proc, path)
    if len(options) < 2:
        raise RuntimeError("Не найдено два допустимых положения чипа")
    centroid = (centroid_info[0], centroid_info[1])
    area = centroid_info[2]
    make_figures(points, triangles, boundary_nodes, proc, centroid, u_minus, u_plus, options)
    write_artifacts(options, valid_count, centroid, area, proc)
    build_report(options, valid_count, centroid, area, proc)
    print(DOCX_OUT)
    print(ARTIFACTS)
    print(options)


if __name__ == "__main__":
    main()
